from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(13)

def h1(text):
    p = doc.add_heading(level=1)
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 70, 127)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p = doc.add_heading(level=2)
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 73, 125)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def h3(text):
    p = doc.add_heading(level=3)
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(54, 95, 145)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(f'→ {text}')
    r.font.name     = 'Times New Roman'
    r.font.size     = Pt(12)
    r.font.italic   = True
    r.font.color.rgb = RGBColor(89, 89, 89)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)

def pb():
    doc.add_page_break()

def center(text, bold=False, size=13, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

# ════════════════════════════════════════════════════
# TRANG BÌA
# ════════════════════════════════════════════════════
center('TRƯỜNG ĐẠI HỌC KINH TẾ – LUẬT', bold=True, size=14)
center('KHOA HỆ THỐNG THÔNG TIN', bold=True, size=13)
doc.add_paragraph()
center('──────────────────────────────')
doc.add_paragraph()
center('BÁO CÁO CUỐI KỲ', bold=True, size=20, color=(0, 70, 127))
center('PHÁT TRIỂN WEB KINH DOANH', bold=True, size=15)
doc.add_paragraph()
center('ĐỀ TÀI:', bold=True)
center('FITFUEL+ — NỀN TẢNG FITNESS TÍCH HỢP', bold=True, size=15, color=(0, 112, 192))
center('THEO DÕI LUYỆN TẬP, ĐẶT ĐỒ ĂN HEALTHY', bold=True, size=15, color=(0, 112, 192))
center('& CHỢ THIẾT BỊ GYM', bold=True, size=15, color=(0, 112, 192))
doc.add_paragraph()
doc.add_paragraph()
for name, sid in [
    ('Nguyễn Văn A', 'K234111XXX'),
    ('Nguyễn Thị B', 'K234111XXX'),
    ('Trần Văn C',   'K234111XXX'),
    ('Lê Thị D',     'K234111XXX'),
    ('Phạm Văn E',   'K234111XXX'),
]:
    center(f'{name}     {sid}')
doc.add_paragraph()
center('Giảng viên: TS. Trần Duy Thanh', bold=True)
center('Môn học: Phát triển Web Kinh Doanh')
doc.add_paragraph()
center('Thành phố Hồ Chí Minh, 2026')
pb()

# ════════════════════════════════════════════════════
# TRANG ĐẦU
# ════════════════════════════════════════════════════
for title, ghi_chu in [
    ('LỜI CẢM ƠN',             'Viết lời cảm ơn tại đây.'),
    ('MỤC LỤC',                'Tạo mục lục tự động: References → Table of Contents trong Word.'),
    ('DANH MỤC TỪ VIẾT TẮT',  'Bảng từ viết tắt: API, JWT, OTP, PR (Personal Record), TDEE, XP, BPMN, DFD, ERD, QR, MVP, FR, NFR, FE, BE, ORM, UI/UX...'),
    ('DANH MỤC HÌNH ẢNH',     'Danh sách hình tự động: References → List of Figures trong Word.'),
    ('DANH MỤC BẢNG',         'Danh sách bảng tự động: References → List of Tables trong Word.'),
]:
    p = doc.add_heading(level=1); p.clear()
    r = p.add_run(title)
    r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    r.font.bold = True; r.font.color.rgb = RGBColor(0, 70, 127)
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    note(ghi_chu)
    pb()

# ════════════════════════════════════════════════════
# CHƯƠNG 1
# ════════════════════════════════════════════════════
h1('CHƯƠNG 1. GIỚI THIỆU')

h2('1.1. Tổng quan dự án')
note('Giới thiệu FitFuel+ là nền tảng web kết hợp 3 mảng: Gym Tracking + Food Order + Gear Hub.')
note('Điểm cốt lõi: dữ liệu 3 mảng "nói chuyện" với nhau — tập nhóm cơ nào, gợi ý đồ ăn đó.')
note('Hai kênh: Customer Portal (Guest & Member) và Back-office Portal (Vendor, Gym Owner, Admin).')
note('Điểm kết nối trung tâm: Fitness Passport — hồ sơ thể hình cá nhân ghi lại toàn bộ hành trình.')

h2('1.2. Đặt vấn đề')
note('Vấn đề 1: Thiếu công cụ theo dõi tiến độ tập luyện — không biết mình có tiến bộ không.')
note('Vấn đề 2: Không biết nên ăn gì sau buổi tập — thông tin dinh dưỡng rời rạc, mâu thuẫn, không kết nối với gym log.')
note('Vấn đề 3: Thiếu thiết bị tập nhưng không muốn mua mới — không có nền tảng thuê gym gear uy tín tại VN.')
note('Vấn đề 4: Thiếu động lực duy trì — tập một mình, không có cơ chế khen thưởng hay cộng đồng.')

h2('1.3. Mục tiêu dự án')

h3('1.3.1. Mục tiêu kinh doanh')
note('Xây dựng hệ sinh thái tạo vòng lặp giá trị: Tập → Gợi ý ăn → Mua/Thuê gear → Tiếp tục tập.')
note('Tạo nguồn thu B2B qua Vendor Portal (quán ăn) và Gym Owner Dashboard (phòng tập).')
note('Tăng gắn kết người dùng qua Fitness Passport, FitCoin và hệ thống Gamification.')

h3('1.3.2. Mục tiêu hệ thống (đo lường được)')
note('Log buổi tập và bài tập trong dưới 60 giây.')
note('Gợi ý món ăn phù hợp dựa trên nhóm cơ vừa tập (rule-based).')
note('Checkout food tối đa 3 bước (cả Guest OTP lẫn Member).')
note('Gear Lifecycle hiển thị toàn bộ lịch sử trong 1 màn hình.')
note('Trang load dưới 2 giây trên mobile.')

h3('1.3.3. Mục tiêu kỹ thuật')
note('REST API chuẩn hóa với JWT authentication và phân quyền theo role.')
note('Responsive design mobile-first từ 375px, tương thích Chrome, Firefox, Safari, Edge.')
note('Continuous Deployment: Frontend → Vercel, Backend → Railway/Render.')
note('Bảo mật: bcrypt 10 rounds, HTTPS, chống SQL injection, XSS, CSRF.')

h2('1.4. Phạm vi dự án')

h3('1.4.1. Trong phạm vi (8 Module)')
note('Module 1 — Quản lý tài khoản: Đăng ký, Đăng nhập, Guest OTP, Fitness Passport.')
note('Module 2 — Gym Tracking: Session, Log bài tập (Set/Rep/Weight), PR, Smart Suggestion, Streak, Check-in QR.')
note('Module 3 — Food Order: Browse, Filter, Cart, Checkout, AI Gợi ý, TDEE, Macro Dashboard, Review có ảnh, Meal Prep.')
note('Module 4 — Gear Hub: Listing, Gear Lifecycle, Thuê, Mua, Đăng bán, QR Code, AI Gear Suggestion.')
note('Module 5 — Gamification: XP, Level, Badge, Streak, Weekly Challenge, Ranking Board.')
note('Module 6 — Social: Feed, Follow, Referral Program.')
note('Module 7 — Payment & FitCoin: VNPay/Momo sandbox, Earn/Spend/Deposit FitCoin, gia hạn membership.')
note('Module 8 — Admin & B2B: Food Vendor Portal, Gym Owner Dashboard, Admin Panel.')

h3('1.4.2. Ngoài phạm vi')
note('Wearable device (Apple Watch, Fitbit). Native mobile app (chỉ responsive web).')
note('Video call / Live streaming với PT. Chat real-time giữa users.')
note('Thanh toán thật (chỉ sandbox). AI/ML phức tạp (chỉ rule-based). Logistics tự động.')

h2('1.5. Thông tin thương hiệu')
note('Tên: FitFuel+. Slogan: "Train Smart. Eat Right. Gear Up."')
note('Logo, màu nhận diện thương hiệu. Chèn hình logo và banner tại đây.')
note('Tầm nhìn (Vision): trở thành hệ sinh thái fitness toàn diện tại Việt Nam cho người 18–35 tuổi ở thành phố lớn.')
note('Sứ mệnh (Mission): kết nối dữ liệu tập luyện — dinh dưỡng — thiết bị, giúp người dùng đạt mục tiêu thể hình bền vững.')

h2('1.6. Phân tích đối thủ cạnh tranh')
note('Bảng so sánh FitFuel+ với MyFitnessPal, Strong/JEFIT, Shopee, GrabFood.')
note('Tiêu chí so sánh: Gym Tracking, Food Order + Macro, AI Gợi ý từ Gym, Gear Marketplace, Gear Lifecycle, Gamification, FitCoin, B2B Dashboard, Guest Checkout, Fitness Passport.')
note('Kết luận: không nền tảng nào kết hợp cả 3 mảng theo cách cho dữ liệu "nói chuyện" với nhau.')

h2('1.7. Cấu trúc báo cáo')
note('Chương 1 — Giới thiệu: bối cảnh, vấn đề, mục tiêu, phạm vi.')
note('Chương 2 — Cơ sở lý thuyết: các công nghệ và framework sử dụng.')
note('Chương 3 — Phân tích và thiết kế hệ thống: FR, NFR, Use Case, BPMN, DFD, CSDL, Kiến trúc, Wireframe, Sequence Diagram.')
note('Chương 4 — Triển khai website: FE/BE/DB, tích hợp, giao diện, kiểm thử, đánh giá.')
note('Chương 5 — Kết luận và hướng phát triển.')
pb()

# ════════════════════════════════════════════════════
# CHƯƠNG 2
# ════════════════════════════════════════════════════
h1('CHƯƠNG 2. CƠ SỞ LÝ THUYẾT')

h2('2.1. HyperText Markup Language (HTML)')
note('Vai trò: nền tảng cấu trúc giao diện Customer Portal và Back-office Portal.')

h2('2.2. CSS & Tailwind CSS')
note('Utility-first CSS framework, responsive design mobile-first (min 375px), hệ thống design nhất quán, JIT compiler giảm bundle size.')

h2('2.3. JavaScript / JSX')
note('Xử lý logic phía client, event handling, JSX là cú pháp mở rộng cho React.')

h2('2.4. React.js')
note('Component-based architecture, Virtual DOM, React Router v6 (Protected Routes), React Hooks (useState, useEffect, useContext).')

h2('2.5. Vite (Build Tool)')
note('Dev server siêu nhanh với native ES modules. Build production tối ưu cho Vercel deployment.')

h2('2.6. Node.js và Express.js')
note('Node.js runtime JavaScript phía server. Express.js framework REST API: routing, middleware pipeline, HTTP utilities. Toàn bộ backend FitFuel+ chạy trên stack này.')

h2('2.7. MySQL / PostgreSQL và RESTful API')
note('Cơ sở dữ liệu quan hệ (ACID compliance, hỗ trợ JOIN phức tạp giữa các module). API chuẩn REST: GET/POST/PUT/PATCH/DELETE, JSON response nhất quán.')

h2('2.8. ORM — Sequelize / Prisma')
note('Ánh xạ object-relational, quản lý migration, bảo vệ SQL injection, type-safe database access.')

h2('2.9. Công cụ phát triển')

h3('2.9.1. Visual Studio Code')
note('Editor chính, tích hợp ESLint, Prettier, GitLens.')

h3('2.9.2. GitHub')
note('Version control. Branching: main (production) / dev (tích hợp) / feature/* (từng tính năng). Pull Request workflow.')

h3('2.9.3. Postman')
note('Kiểm thử và tài liệu hóa API theo từng module (Auth, Gym, Food, Gear, Order, Payment, Admin).')

h3('2.9.4. Draw.io / Figma')
note('Draw.io: vẽ sơ đồ hệ thống (Use Case, BPMN, DFD, ERD, Sequence Diagram, Architecture). Figma: wireframe giao diện.')

h3('2.9.5. Cloudinary')
note('Cloud storage ảnh: ảnh đánh giá món ăn (Review) và ảnh thiết bị Gear. SDK tích hợp Node.js.')

h2('2.10. Áp dụng lý thuyết vào dự án')
note('Bảng tổng hợp: Công nghệ | Áp dụng cụ thể vào module nào trong FitFuel+.')
pb()

# ════════════════════════════════════════════════════
# CHƯƠNG 3
# ════════════════════════════════════════════════════
h1('CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG')

h2('3.1. Yêu cầu hệ thống')

h3('3.1.1. Yêu cầu chức năng (FR)')
note('Bảng FR đầy đủ FR-001 → FR-050, phân nhóm theo 8 module.')
note('Cột: ID | Mô tả yêu cầu | Độ ưu tiên (Cao / Trung bình / Thấp).')
note('Module 1 — Quản lý tài khoản (FR-001 → FR-006): Đăng ký, Đăng nhập, Guest OTP, Merge Guest, Cập nhật profile, Fitness Passport.')
note('Module 2 — Gym Tracking (FR-007 → FR-014): Tạo session, Log exercise, Tính PR, Progress chart, Lịch sử, Smart Suggestion, Check-in QR, Thống kê.')
note('Module 3 — Food Order (FR-015 → FR-026): Danh sách + filter, Chi tiết, Add to cart từ trang chủ, Sửa thuộc tính trong cart, Checkout Member, Checkout Guest OTP, Re-order, AI Gợi ý, TDEE, Macro Dashboard, Đánh giá có ảnh, Meal Prep.')
note('Module 4 — Gear Hub (FR-027 → FR-033): Listing, Gear Lifecycle, Thuê, Mua, Đăng bán, QR Code, AI Gear Suggestion.')
note('Module 5 — Gamification (FR-034 → FR-038): XP + Level, Badge, Streak, Weekly Challenge, Ranking Board.')
note('Module 6 — Social (FR-042 → FR-044): Post milestone lên Feed, Follow, Referral Program.')
note('Module 7 — Payment & FitCoin (FR-039 → FR-041): Thanh toán sandbox, Quản lý FitCoin, Gia hạn membership.')
note('Module 8 — Admin & B2B (FR-045 → FR-050): Vendor đăng SP, Vendor xử lý đơn, Vendor analytics, Gym quản lý member, Gym thông báo, Admin duyệt, Admin tranh chấp, Notification.')

h3('3.1.2. Yêu cầu phi chức năng (NFR)')
note('Bảng NFR theo 5 nhóm. Cột: ID | Mô tả | Chỉ số đo.')
note('Hiệu năng: Load trang < 2s (mobile), API response < 500ms, chịu 100 concurrent users.')
note('Bảo mật: bcrypt 10 rounds, HTTPS, JWT 7 ngày, chống SQL injection / XSS / CSRF, validate cả client lẫn server.')
note('Khả dụng: Responsive từ 375px, tương thích 4 browser, checkout ≤ 3 bước, font ≥ 14px, contrast ≥ 4.5:1, tap target ≥ 44×44px.')
note('Tin cậy: Backup hàng ngày, uptime 99% (demo), lỗi hiển thị thân thiện không lộ stack trace.')
note('Quyền riêng tư: Toggle public/private Fitness Passport, toggle từng body photo, quyền xóa tài khoản.')

h2('3.2. Phân tích tác nhân và Use Case')

h3('3.2.1. Danh sách tác nhân')
note('Bảng 8 tác nhân. Cột: Tên | Loại (Chính/Phụ) | Mô tả | Tương tác chính | Không được phép.')
note('Guest, Member, Food Vendor, Gear Seller, Gym Owner, Admin, Timer (hệ thống định thời tự động), Payment Gateway (VNPay/Momo).')

h3('3.2.2. Biểu đồ Use Case tổng thể')
note('Chèn biểu đồ Use Case tổng thể (UC-01 → UC-56). System Boundary bao gồm 8 module. Vẽ bằng Draw.io.')

h3('3.2.3. Biểu đồ Use Case chi tiết theo module')
note('Chèn biểu đồ Use Case chi tiết Module Food Order (với quan hệ <<include>>, <<extend>>).')
note('Chèn biểu đồ Use Case chi tiết Module Gear Hub.')

h3('3.2.4. Đặc tả Use Case')
note('Đặc tả ~10 UC quan trọng nhất. Mỗi UC gồm: ID | Tên | Tác nhân | Mục tiêu | Điều kiện tiên | Điều kiện sau | Luồng cơ bản | Luồng thay thế | Luồng ngoại lệ | Quy tắc nghiệp vụ.')
note('Danh sách UC cần đặc tả: UC-01 Đăng ký, UC-03 Guest OTP, UC-08 Log Exercise, UC-09 Tính PR, UC-21 Checkout, UC-24 AI Gợi ý thực đơn, UC-28 Xem Gear Lifecycle, UC-29 Thuê Gear, UC-33 Đăng bán Gear, UC-42 Thanh toán.')

h2('3.3. Mô hình quy trình nghiệp vụ BPMN')
note('Chèn các sơ đồ BPMN (vẽ bằng Draw.io):')
note('BPMN 3.3.1 — Đăng ký / Đăng nhập / Guest OTP.')
note('BPMN 3.3.2 — Tạo Workout Session & Log Exercise (bao gồm kiểm tra PR, cập nhật Streak).')
note('BPMN 3.3.3 — Đặt đồ ăn: luồng Member 3 bước và luồng Guest OTP (2 lane song song).')
note('BPMN 3.3.4 — AI Gợi ý thực đơn: Nhóm cơ → Ưu tiên Macro → Lọc món ăn → Hiện popup.')
note('BPMN 3.3.5 — Thuê / Mua thiết bị Gear (bao gồm đặt cọc, cập nhật Gear Lifecycle).')
note('BPMN 3.3.6 — Đăng bán/cho thuê Gear → Sinh Gear ID → Tạo QR Code → Tạo Gear Lifecycle lần đầu.')
note('BPMN 3.3.7 — Xử lý thanh toán: VNPay/Momo sandbox callback và nhánh FitCoin.')

h2('3.4. Sơ đồ luồng dữ liệu (DFD)')
note('Chèn DFD Level 0 (Context Diagram): hệ thống FitFuel+ và các external entities (Member, Guest, Vendor, Gym Owner, Admin, Payment Gateway).')
note('Chèn DFD Level 1: luồng dữ liệu chi tiết Customer Subsystem (Gym Tracking → AI Suggestion → Food Order → Gear Hub).')

h2('3.5. Thiết kế cơ sở dữ liệu')

h3('3.5.1. Sơ đồ thực thể - quan hệ (ERD)')
note('Chèn ERD đầy đủ tại đây. Vẽ bằng Draw.io hoặc dbdiagram.io.')
note('Các thực thể chính: USERS, FITNESS_PASSPORT, WORKOUT_SESSIONS, SESSION_EXERCISES, EXERCISES, FOOD_ITEMS, ORDERS, ORDER_ITEMS, GEARS, GEAR_LIFECYCLE, BADGES, USER_BADGES, XP_LOG, FITCOIN_LOG, VENDORS, GYM_OWNERS, REVIEWS, SOCIAL_POSTS, FOLLOWS.')

h3('3.5.2. Lược đồ cơ sở dữ liệu (Database Schema)')
note('Bảng định nghĩa từng table: Tên cột | Kiểu dữ liệu | Ràng buộc (PK/FK/NOT NULL/UNIQUE) | Mô tả.')
note('Hoặc chèn các câu lệnh SQL CREATE TABLE.')

h3('3.5.3. Mô hình dữ liệu Rule-based AI Mapping')
note('Bảng MUSCLE_MACRO_MAPPING: Nhóm cơ (Ngực/Lưng/Chân/Vai/Tay/Core/Toàn thân) | Ưu tiên Macro | Tags lọc món ăn.')
note('Fallback logic: nếu không có gym log hôm nay → gợi ý theo user.goal (bulk/cut/maintain).')

h2('3.6. Kiến trúc hệ thống')

h3('3.6.1. Kiến trúc tổng thể')
note('Chèn sơ đồ kiến trúc 3 tầng tại đây.')
note('Presentation: React/Vite → Vercel CDN.')
note('Application: Express.js → Railway/Render.')
note('Data: MySQL → Aiven (free tier). Media: Cloudinary. Payment: VNPay/Momo sandbox.')

h3('3.6.2. Kiến trúc module Rule-based AI Recommendation')
note('Pipeline: Kết thúc buổi tập → Đọc muscle_group từ session → Tra bảng mapping → Query FOOD_ITEMS theo tags → Trả về top 3 → Hiện popup.')
note('Fallback: không có session hôm nay → đọc user.goal → áp dụng macro filter tương ứng.')

h3('3.6.3. Kiến trúc Gear Lifecycle & QR Code')
note('Luồng: Đăng gear → Sinh UUID Gear ID → Upload ảnh Cloudinary → Generate QR code → Upload QR lên Cloudinary → Tạo GEAR_LIFECYCLE entry đầu tiên (action = listed).')
note('Cơ chế: GEAR_LIFECYCLE là append-only — mỗi sự kiện (thuê/trả/bán) tạo entry mới, không sửa entry cũ.')

h2('3.7. Thiết kế giao diện')

h3('3.7.1. Sitemap')
note('Chèn Sitemap kênh người dùng (Customer Portal): Home / Auth / Gym / Food / Cart / Checkout / Gear / Fitness Passport / Social / FitCoin.')
note('Chèn Sitemap kênh quản trị (Back-office): Vendor Portal / Gym Owner Dashboard / Admin Panel.')

h3('3.7.2. Wireframe')
note('Chèn wireframe (low-fidelity) cho từng màn hình chính:')
note('Đăng ký / Đăng nhập / Quên mật khẩu.')
note('Trang chủ: Hero + Food card grid (nút [+] add to cart nhanh) + Gear nổi bật.')
note('Gym Tracking: Chọn nhóm cơ → Form log Set/Rep/Weight động → Hiển thị PR badge inline.')
note('Popup AI Gợi ý thực đơn (overlay sau khi bấm Kết thúc buổi tập).')
note('Danh sách Food (filter sidebar) + Chi tiết Food (macro breakdown donut chart + review).')
note('Giỏ hàng (sửa qty/size/topping trực tiếp) + Checkout 3 bước.')
note('Gear Hub: Listing grid + Gear Lifecycle timeline dọc + Form đăng bán gear.')
note('Fitness Passport: Stats row + Badge grid + Body photo gallery.')
note('Food Vendor Portal: Form thêm sản phẩm + Kanban quản lý đơn + Analytics.')
note('Gym Owner Dashboard + Admin Panel.')

h3('3.7.3. Luồng xử lý request Backend')
note('Chèn sơ đồ luồng: Client → Express Router → JWT Middleware → Role Middleware → Validate Middleware → Controller → Service → Model (ORM) → Database → JSON Response.')

h2('3.8. Sơ đồ tuần tự (Sequence Diagram)')
note('Chèn các Sequence Diagram sau (vẽ bằng Draw.io hoặc PlantUML):')
note('SD-01 — Luồng Đăng nhập và JWT.')
note('SD-02 — Luồng Guest OTP Checkout.')
note('SD-03 — Luồng Tạo Workout Session + Log Exercise + Kiểm tra PR.')
note('SD-04 — Luồng AI Gợi ý thực đơn sau buổi tập.')
note('SD-05 — Luồng Đăng bán Gear → Sinh Gear ID → QR Code → Gear Lifecycle.')
pb()

# ════════════════════════════════════════════════════
# CHƯƠNG 4
# ════════════════════════════════════════════════════
h1('CHƯƠNG 4. TRIỂN KHAI WEBSITE')

h2('4.1. Triển khai Frontend')
note('Mô tả cấu trúc thư mục React/Vite: /components, /pages, /hooks, /store, /api, /utils, /router.')
note('State management: Zustand (cart, auth, notification). Routing: React Router v6 với ProtectedRoute và GuestRoute.')
note('API calls: Axios instance với JWT interceptor tự động gắn Authorization header.')

h2('4.2. Triển khai Backend')
note('Cấu trúc Express.js: /routes, /controllers, /services, /middlewares, /models, /config.')
note('JWT: jsonwebtoken, 7 ngày, RS256/HS256. OTP: 6 chữ số, cache TTL 5 phút, khóa sau 3 lần sai 15 phút.')
note('bcrypt: 10 salt rounds. Input validation: express-validator. Error handler: global middleware trả JSON chuẩn.')

h2('4.3. Triển khai cơ sở dữ liệu')
note('MySQL schema với Sequelize migrations. Foreign key ON DELETE CASCADE khi phù hợp.')
note('GEAR_LIFECYCLE: append-only, không cho phép UPDATE entry cũ.')
note('Seed data demo: 30+ món ăn (đầy đủ macro + tags), 20+ thiết bị gear (kèm 3–5 Lifecycle entries mỗi cái), 50+ bài tập theo 7 nhóm cơ, tài khoản demo (Guest phone, 2 Members, 1 Vendor, 1 Gym Owner, 1 Admin).')

h2('4.4. Tích hợp bổ sung')

h3('4.4.1. Module AI Gợi ý thực đơn (Rule-based)')
note('Hàm recommendFood(session_id): đọc muscle_group → tra mapping object → query FOOD_ITEMS → trả top 3 theo protein density.')
note('Fallback recommendByGoal(user_id): đọc user.goal → áp macro filter → query → trả top 3.')
note('Không cần thư viện ML. Response time < 80ms (1 DB query với indexed column).')

h3('4.4.2. Tích hợp cổng thanh toán VNPay / Momo Sandbox')
note('Tạo URL thanh toán với chữ ký HMAC-SHA512 → Redirect sang gateway → Nhận callback → Xác minh chữ ký → Cập nhật trạng thái đơn → Thông báo Vendor.')
note('Nhánh FitCoin: trừ số dư trong DB transaction để tránh race condition.')

h3('4.4.3. Sinh Gear ID & QR Code')
note('Sinh UUID v4 làm gear_id. Dùng thư viện qrcode tạo QR image buffer → upload Cloudinary → lưu secure_url.')
note('Frontend dùng html5-qrcode quét QR qua camera trên Vercel HTTPS (camera cần HTTPS permission).')

h3('4.4.4. Upload ảnh Cloudinary')
note('multer middleware stream multipart/form-data trực tiếp lên Cloudinary (không lưu local disk).')
note('Giới hạn: 5MB/file, JPG/PNG/WebP. Review food: 1 ảnh. Gear listing: 2–8 ảnh.')

h2('4.5. Giao diện chức năng chính')
note('Với mỗi màn hình: chèn ảnh chụp màn hình thực tế + mô tả ngắn chức năng.')

h3('4.5.1. Thành phần dùng chung (General Components)')
note('Navbar (logo, nav links, cart badge, avatar dropdown). Footer. Toast notification. Loading skeleton. Protected Route wrapper.')

h3('4.5.2. Giao diện kênh người dùng (Customer Interface)')
note('Đăng ký / Đăng nhập / Quên mật khẩu.')
note('Trang chủ.')
note('Gym Tracking: Danh sách session / Tạo session / Form log bài tập / PR badge / Thống kê.')
note('Popup AI Gợi ý thực đơn (sau khi kết thúc buổi tập).')
note('Food: Danh sách (filter) / Chi tiết (macro breakdown) / Đánh giá có ảnh.')
note('Giỏ hàng / Checkout 3 bước (Member) / Checkout OTP (Guest) / Re-order / Meal Prep.')
note('Gear Hub: Listing / Gear Lifecycle timeline / Thuê / Mua / Form đăng bán / Quét QR.')
note('Fitness Passport (Stats + Badge + Body photos).')
note('Social Feed / Follow / Xem Fitness Passport người khác.')
note('Ví FitCoin (số dư + lịch sử giao dịch).')

h3('4.5.3. Giao diện Admin')
note('Admin Dashboard / Duyệt Vendor & Gym Owner / Xử lý tranh chấp / Báo cáo tổng thể.')

h3('4.5.4. Giao diện B2B')
note('Food Vendor Portal: Quản lý sản phẩm / Kanban đơn hàng (Pending → Confirmed → Delivered) / Analytics (doanh thu, top món, review).')
note('Gym Owner Dashboard: Danh sách member / Gửi thông báo.')

h2('4.6. Kiểm thử chức năng')
note('Bảng test case các luồng cốt lõi. Cột: TC ID | Module | Kịch bản | Input | Kết quả mong đợi | Kết quả thực tế.')
note('Tối thiểu 15–20 test case bao phủ FR ưu tiên Cao: Auth, Guest OTP (happy path + lock sau 3 lần), Gym (log + PR), Food (add to cart từ homepage, sửa cart, checkout), AI Gợi ý, Gear Lifecycle, Payment, Vendor (thêm sản phẩm, xác nhận đơn).')

h2('4.7. Đánh giá hệ thống')

h3('4.7.1. Mục tiêu đánh giá')
note('Xác nhận FR ưu tiên Cao đã triển khai đầy đủ và hoạt động đúng trên môi trường production.')
note('Kiểm tra NFR về hiệu năng, bảo mật, khả dụng.')

h3('4.7.2. Phạm vi đánh giá')
note('Customer Portal (Guest + Member), Back-office Portal (Vendor, Gym Owner, Admin), toàn bộ API endpoints.')

h3('4.7.3. Cấu hình môi trường kiểm thử')
note('Bảng: Frontend URL (Vercel) / Backend URL (Railway) / Database (Aiven MySQL) / Thiết bị (Desktop Chrome, Mobile Safari, Android Chrome) / Công cụ (Lighthouse, Postman).')

h3('4.7.4. Kiểm thử API bằng Postman')
note('Mô tả Postman Collection theo folder: Auth / Gym / Food / Recommendation / Gear / Payment / Admin & B2B.')
note('Kịch bản test: Happy path, Edge case (OTP lock, PR detection), Security (JWT hết hạn, SQL injection input, XSS payload).')

h3('4.7.5. Kết quả kiểm thử chức năng')
note('Bảng tổng hợp: Module | Tổng FR | Đã hoàn thành | Tỷ lệ (%) | Ghi chú. Tổng thể khoảng 86%.')

h3('4.7.6. Phân tích yêu cầu phi chức năng')
note('Bảng: Tiêu chí NFR | Phương pháp đo | Kết quả thực tế | Đạt / Không đạt.')
note('Bao gồm: Page load (Lighthouse Mobile), API response (Postman average), Responsive 375px, Checkout 3 bước, bcrypt verify, JWT expiry 401, XSS input bị sanitize.')

h3('4.7.7. Ràng buộc và đánh giá rủi ro')
note('Bảng: Rủi ro | Mức độ ảnh hưởng | Biện pháp giảm thiểu.')
note('Backend cold start (Railway free tier sleep 15 phút). OTP qua console thay vì SMS thật. Cloudinary bandwidth giới hạn. Sandbox VNPay/Momo không ổn định.')
pb()

# ════════════════════════════════════════════════════
# CHƯƠNG 5
# ════════════════════════════════════════════════════
h1('CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN')

h2('5.1. Kết luận')
note('Tóm tắt kết quả đạt được: tỷ lệ hoàn thành FR (~86%), vòng lặp giá trị Gym → AI Food → Gear, Gear Lifecycle minh bạch, Guest OTP checkout, kiến trúc B2B đa vai trò.')
note('Xác nhận đạt các NFR: hiệu năng, bảo mật, khả dụng trên môi trường production.')
note('Điểm nổi bật: FitFuel+ là nền tảng duy nhất tại VN kết hợp cả 3 mảng theo cách dữ liệu "nói chuyện" với nhau.')

h2('5.2. Hạn chế')

h3('5.2.1. Cổng thanh toán ở chế độ sandbox')
note('VNPay/Momo chỉ sandbox, không xử lý tiền thật. Sandbox đôi khi không ổn định trong demo.')

h3('5.2.2. AI rule-based thay vì Machine Learning thực sự')
note('Mapping tĩnh không cá nhân hóa theo lịch sử từng user. Chưa có collaborative filtering hay content-based recommendation.')

h3('5.2.3. OTP qua console thay vì SMS thật')
note('Chưa tích hợp SMS gateway thật (Twilio, ESMS VN) do ràng buộc ngân sách bằng 0.')

h3('5.2.4. Ràng buộc free-tier deployment')
note('Backend sleep sau 15 phút idle → cold start ~5 giây. Aiven MySQL free tier không có replication.')

h2('5.3. Hướng phát triển')

h3('5.3.1. Ứng dụng mobile native')
note('React Native: tái sử dụng backend API, thêm push notification, QR scan native, offline mode cho gym log.')

h3('5.3.2. Tích hợp thiết bị wearable')
note('Apple Health / Google Fit API: tự động import workout data, bỏ bước nhập tay.')

h3('5.3.3. AI/ML Recommendation thực sự')
note('Collaborative filtering (lịch sử tương tác). Content-based filtering (macro profile cá nhân). Hybrid model cho cold-start users.')

h3('5.3.4. Production deployment & Containerization')
note('Docker + Kubernetes cho backend. Managed DB cluster (AWS RDS). Global CDN cho frontend.')

h3('5.3.5. Hoàn thiện hệ thống Gamification')
note('Weekly Challenge FE (backend đã xong). Ranking Board real-time. FitCoin economy cân bằng với cơ chế chống lạm phát.')

h3('5.3.6. Tính năng real-time')
note('WebSocket (Socket.io): thông báo PR real-time, cập nhật trạng thái đơn live. WebRTC: video call với Personal Trainer.')
pb()

# ════════════════════════════════════════════════════
# TÀI LIỆU THAM KHẢO
# ════════════════════════════════════════════════════
p = doc.add_heading(level=1); p.clear()
r = p.add_run('TÀI LIỆU THAM KHẢO')
r.font.name='Times New Roman'; r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=RGBColor(0,70,127)
p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(4)
note('Liệt kê tài liệu tham khảo theo định dạng APA hoặc quy định của trường.')
note('Bao gồm: tài liệu công nghệ (React, Node.js, Express, MySQL...), nghiên cứu thị trường fitness VN, tài liệu tích hợp (VNPay, Cloudinary, Momo), OWASP Security...')
pb()

# ════════════════════════════════════════════════════
# PHÂN CÔNG NHIỆM VỤ
# ════════════════════════════════════════════════════
p = doc.add_heading(level=1); p.clear()
r = p.add_run('PHÂN CÔNG NHIỆM VỤ NHÓM')
r.font.name='Times New Roman'; r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=RGBColor(0,70,127)
p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(4)

h2('Giai đoạn 1 — Phân tích & Thiết kế')
note('Bảng phân công. Cột: Thành viên | Nhiệm vụ | % đóng góp.')

h2('Giai đoạn 2 — Lập trình')
note('Bảng phân công. Cột: Thành viên | Module phụ trách | % đóng góp.')
note('TV1 (FE Lead): Base React/Tailwind, Gym Tracking UI, B2B Dashboard UI.')
note('TV2 (FE): Food Order UI (List, Detail, Cart, Checkout, Meal Prep, Review có ảnh).')
note('TV3 (FE/Tích hợp): Gear Hub UI (Lifecycle Timeline, QR Code scan), Social & Community UI.')
note('TV4 (BE Lead): Base Express.js, JWT Auth, Gym Tracking API, Rule-based AI API, Social API.')
note('TV5 (BE/Database): Schema & Migrations, Food/Gear/Review API, Analytics API, Seed Data.')

h2('Giai đoạn 3 — Hoàn thiện báo cáo')
note('Bảng phân công. Cột: Thành viên | Phần báo cáo | % đóng góp.')

h2('Bảng đánh giá thành viên')
note('Bảng: Thành viên | Tự đánh giá (%) | Đánh giá chéo trung bình (%) | Điểm cuối.')

# ════════════════════════════════════════════════════
# LƯU FILE
# ════════════════════════════════════════════════════
output = r'D:\doanWEDKD\FitFuel+_Outline.docx'
doc.save(output)
print(f'[OK] Da luu: {output}')
