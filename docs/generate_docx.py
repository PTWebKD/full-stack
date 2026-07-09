import os
import re
import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Colors for styling
COLOR_PRIMARY_HEX = 'FF5722'  # FitFuel+ Orange
COLOR_GRAY_HEX = '71717A'
COLOR_LIGHT_BG_HEX = 'F4F4F5'
COLOR_BORDER_HEX = 'D4D4D8'

def set_cell_background(cell, fill_hex):
    """Set background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'w:top': top, 'w:bottom': bottom, 'w:left': left, 'w:right': right}
    for key, border_style in borders.items():
        if border_style is not None:
            node = OxmlElement(key)
            node.set(qn('w:val'), border_style.get('val', 'single'))
            node.set(qn('w:sz'), str(border_style.get('sz', 4)))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_markdown_paragraph(doc, text, style='Normal', is_bullet=False):
    """Adds a paragraph with inline bold and italic parsing."""
    if is_bullet:
        # Strip list prefix
        text = re.sub(r'^[\-\*\+]\s+', '', text)
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph(style=style)
        
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    if is_bullet:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.75)
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Clean text of emojis if desired, but let's keep it clean
    text = text.replace('⭐⭐⭐⭐⭐', '★★★★★')
    text = text.replace('⭐⭐⭐⭐', '★★★★☆')
    text = text.replace('⭐⭐⭐', '★★★☆☆')
    text = text.replace('⭐⭐', '★★☆☆☆')
    text = text.replace('⭐', '★☆☆☆☆')
    text = text.replace('❌', '—')

    # Bold parser
    parts = text.split('**')
    is_bold = False
    for part in parts:
        # Italic parser
        subparts = part.split('*')
        is_italic = False
        for subpart in subparts:
            if subpart:
                run = p.add_run(subpart)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                if is_bold:
                    run.bold = True
                if is_italic:
                    run.italic = True
            is_italic = not is_italic
        is_bold = not is_bold
    return p

def create_document():
    # Read Markdown
    md_path = 'd:/doanWEDKD/docs/CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI (REVISED).md'
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()

    # Set margins (Academic standard in Vietnam: Top 2.5, Bottom 2.5, Left 3.0, Right 2.0)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # Set normal style font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Parse state variables
    in_table = False
    table_lines = []
    fig_count = 1

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines if not in table
        if not line:
            if in_table:
                # Process the table we finished
                process_table(doc, table_lines)
                table_lines = []
                in_table = False
            i += 1
            continue

        # If it's a table row
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # Table ended
            process_table(doc, table_lines)
            table_lines = []
            in_table = False

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(line[2:].upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.bold = True
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
            run.italic = True
        # Horizontal Rule
        elif line.startswith('---'):
            # Ignore or add page break if Chapter transition, but here we just ignore
            pass
        # Image
        elif line.startswith('!['):
            # Parse image: ![Alt Text](Path)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                img_rel_path = match.group(2)
                # Convert relative path to absolute
                img_abs_path = os.path.normpath(os.path.join('d:/doanWEDKD/docs', img_rel_path))
                
                if os.path.exists(img_abs_path):
                    # Insert Image
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(6)
                    run_img = p_img.add_run()
                    run_img.add_picture(img_abs_path, width=Cm(15))  # Fits nicely within standard margin width (16cm printable)
                    
                    # Caption
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run(f"Hình 1.{fig_count}: {alt_text}")
                    run_cap.font.name = 'Times New Roman'
                    run_cap.font.size = Pt(11)
                    run_cap.italic = True
                    fig_count += 1
                else:
                    print(f"Warning: Image file {img_abs_path} does not exist.")
        # Lists (bullet points)
        elif line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s', line):
            # We treat them all as bullet list items
            add_markdown_paragraph(doc, line, is_bullet=True)
        # Citation lines or small source text
        elif line.startswith('*Nguồn:') or line.startswith('*Thực tiễn:'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Strip markdown italics and run as normal italic
            clean_text = line.replace('*', '')
            run = p.add_run(clean_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.italic = True
        # Plain text paragraphs
        else:
            add_markdown_paragraph(doc, line)
            
        i += 1

    # Save Document
    out_docx_path = 'd:/doanWEDKD/docs/CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI (REVISED).docx'
    try:
        doc.save(out_docx_path)
        print("Document successfully created.")
    except PermissionError:
        try:
            out_docx_path_alt = 'd:/doanWEDKD/docs/CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI (REVISED_NEW).docx'
            doc.save(out_docx_path_alt)
            print("Document successfully created (saved as REVISED_NEW).")
        except PermissionError:
            out_docx_path_alt2 = 'd:/doanWEDKD/docs/CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI (REVISED_LATEST).docx'
            doc.save(out_docx_path_alt2)
            print("Document successfully created (saved as REVISED_LATEST).")

def process_table(doc, lines):
    """Helper to convert Markdown table lines to Docx Table."""
    if len(lines) < 2:
        return
        
    # Parse headers and data
    headers = [c.strip() for c in lines[0].split('|')[1:-1]]
    data_rows = []
    
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Apply ratings translation
        cells = [c.replace('⭐⭐⭐⭐⭐', '★★★★★').replace('⭐⭐⭐⭐', '★★★★☆')
                  .replace('⭐⭐⭐', '★★★☆☆').replace('⭐⭐', '★★☆☆☆')
                  .replace('⭐', '★☆☆☆☆').replace('❌', '—') for c in cells]
        data_rows.append(cells)
        
    cols_count = len(headers)
    rows_count = len(data_rows) + 1
    
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    # Border specifications
    border_style = {'val': 'single', 'sz': 4, 'color': COLOR_BORDER_HEX}
    
    # Write Headers
    hdr_cells = table.rows[0].cells
    for col_idx in range(cols_count):
        hdr_cells[col_idx].text = headers[col_idx]
        set_cell_background(hdr_cells[col_idx], COLOR_LIGHT_BG_HEX)
        set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=150, right=150)
        set_cell_borders(hdr_cells[col_idx], top=border_style, bottom=border_style, left=border_style, right=border_style)
        
        # Style formatting
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            
    # Write Data
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx in range(min(cols_count, len(row_data))):
            row_cells[col_idx].text = row_data[col_idx]
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            set_cell_borders(row_cells[col_idx], top=border_style, bottom=border_style, left=border_style, right=border_style)
            
            p = row_cells[col_idx].paragraphs[0]
            # Left align first col, center others
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            
            # Simple markdown bold parser inside cells
            cell_text = row_data[col_idx]
            p.text = ""  # Clear and rebuild with runs
            
            parts = cell_text.split('**')
            is_bold = False
            for part in parts:
                if part:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    if is_bold:
                        run.bold = True
                is_bold = not is_bold

    # Add space after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)
    p_after.text = "" # just empty space

if __name__ == '__main__':
    create_document()
