from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ── Default font ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def set_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line)

def add_heading1(text):
    p = doc.add_paragraph()
    para_spacing(p, before=12, after=6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(0, 70, 127))
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    para_spacing(p, before=10, after=4)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(31, 73, 125))
    return p

def add_body(text, indent=True, bold_parts=None):
    """bold_parts: list of substrings to bold"""
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=6, line=22)
    p.paragraph_format.first_line_indent = Cm(1.27) if indent else Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_parts is None:
        run = p.add_run(text)
        set_font(run)
    else:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            before = remaining[:idx]
            if before:
                r = p.add_run(before)
                set_font(r)
            r = p.add_run(bp)
            set_font(r, bold=True)
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            set_font(r)
    return p

def add_note(text):
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=4)
    p.paragraph_format.left_indent  = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=12, italic=True, color=(89, 89, 89))
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(title_p, before=0, after=4)
r = title_p.add_run('PHÂN TÍCH CÁC BÊN LIÊN QUAN')
set_font(r, size=16, bold=True, color=(0, 70, 127))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(sub_p, before=0, after=2)
r = sub_p.add_run('Stakeholder Analysis')
set_font(r, size=13, italic=True, color=(89, 89, 89))

proj_p = doc.add_paragraph()
proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(proj_p, before=0, after=16)
r = proj_p.add_run('Dự án: FitFuel+  |  Môn học: Web Kinh Doanh  |  Ngày: 16/05/2026')
set_font(r, size=12, color=(89, 89, 89))

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
# 1.5  TỔNG QUAN
# ══════════════════════════════════════════════════════════════════════════════
add_heading1('1.5. Các bên liên quan và nhóm người dùng')

add_body(
    'Hệ thống FitFuel+ có nhiều bên liên quan với vai trò và mức độ tương tác khác nhau, '
    'tham gia vào hệ sinh thái kết hợp giữa theo dõi tập luyện, đặt thức ăn healthy và '
    'thuê/mua thiết bị gym. Mỗi nhóm mang lại giá trị riêng biệt và đồng thời phụ thuộc '
    'vào các nhóm khác để duy trì hoạt động của nền tảng.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 1.5.1
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('1.5.1. Nhóm người dùng trực tiếp')

add_body(
    'Nhóm người dùng trực tiếp bao gồm Khách vãng lai (Guest), Thành viên (Member), '
    'Chủ quán ăn healthy (Food Vendor), Người bán/cho thuê thiết bị (Gear Seller) và '
    'Chủ phòng tập (Gym Owner) — là các đối tượng tương tác thường xuyên nhất với hệ '
    'thống và phát sinh phần lớn các nghiệp vụ cốt lõi.',
    bold_parts=[
        'Khách vãng lai (Guest)', 'Thành viên (Member)',
        'Chủ quán ăn healthy (Food Vendor)', 'Người bán/cho thuê thiết bị (Gear Seller)',
        'Chủ phòng tập (Gym Owner)'
    ]
)

add_body(
    'Khách vãng lai (Guest) là người dùng chưa đăng ký tài khoản, có thể xem danh sách thức ăn '
    'và thiết bị, thêm vào giỏ hàng và thực hiện thanh toán thông qua xác thực OTP mà không cần '
    'tạo tài khoản. Đây là nhóm hướng đến trải nghiệm mua hàng nhanh, giảm rào cản gia nhập ban đầu.',
    bold_parts=['Khách vãng lai (Guest)']
)

add_body(
    'Thành viên (Member) là người dùng đã đăng ký tài khoản, được cấp đầy đủ quyền truy cập vào '
    'toàn bộ tính năng: ghi nhật ký tập luyện, đặt thức ăn với AI gợi ý, tham gia Gear Hub, tích '
    'lũy FitCoin, tham gia gamification và mạng xã hội nội bộ. Đây là nhóm người dùng cốt lõi, '
    'mang lại giá trị lâu dài cho nền tảng thông qua dữ liệu hành vi và giao dịch liên tục.',
    bold_parts=['Thành viên (Member)']
)

add_body(
    'Chủ quán ăn healthy (Food Vendor) đăng ký trở thành đối tác cung cấp thực phẩm trên nền '
    'tảng. Nhóm này sử dụng Food Vendor Portal để đăng sản phẩm kèm thông tin macro chi tiết, '
    'tiếp nhận và xử lý đơn hàng, theo dõi doanh thu và phân tích hành vi khách hàng.',
    bold_parts=['Chủ quán ăn healthy (Food Vendor)']
)

add_body(
    'Người đăng bán/cho thuê thiết bị (Gear Seller/Renter) là vai trò phụ khi Gym Owner (được bán/cho thuê) '
    'hoặc Member (chỉ được cho thuê) đăng thiết bị gym lên Gear Hub kèm Gear ID, hình ảnh và lịch sử '
    'sử dụng minh bạch. Vai trò này tạo nên thị trường thiết bị đáng tin cậy, phân biệt FitFuel+ với '
    'các nền tảng mua bán thông thường.',
    bold_parts=['Người đăng bán/cho thuê thiết bị (Gear Seller/Renter)']
)

add_body(
    'Chủ phòng tập kiêm Quản trị viên (Gym Owner) đăng ký tài khoản Business, sử dụng dashboard riêng '
    'để quản lý phòng tập (thành viên, membership) đồng thời nắm toàn quyền quản trị hệ thống FitFuel+ '
    '(duyệt vendor, xử lý tranh chấp, quản lý FitCoin). Nhóm này đại diện cho kênh B2B quan trọng '
    'và là cấp quản lý cao nhất của nền tảng.',
    bold_parts=['Chủ phòng tập kiêm Quản trị viên (Gym Owner)']
)

# ══════════════════════════════════════════════════════════════════════════════
# Summary table 1.5.1
# ══════════════════════════════════════════════════════════════════════════════
add_note('Bảng 1.5.1 – Tổng hợp nhóm người dùng trực tiếp')

tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Tác nhân', 'Đặc điểm', 'Tương tác chính với hệ thống']
rows_data = [
    ['Guest',       'Chưa đăng ký tài khoản',              'Xem food/gear, thêm giỏ hàng, checkout bằng OTP'],
    ['Member',      'Có tài khoản, Fitness Passport, FitCoin', 'Gym tracking, đặt food, Gear Hub, gamification, social'],
    ['Food Vendor', 'Đối tác kinh doanh, có portal riêng', 'Đăng sản phẩm, xử lý đơn hàng, xem analytics'],
    ['Gear Seller', 'Role phụ của Gym Owner/Member',        'Đăng gear, tạo Gear ID, quản lý listing'],
    ['Gym Owner',   'Quản lý hệ thống & phòng tập',        'Quản lý member, duyệt vendor, xử lý tranh chấp, hệ thống FitCoin'],
]

col_widths = [Cm(3.5), Cm(5.5), Cm(8.0)]
for i, cell in enumerate(tbl.rows[0].cells):
    cell.width = col_widths[i]
    shade_cell(cell, '1F497D')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(headers[i])
    set_font(run, size=12, bold=True, color=(255, 255, 255))

for ri, row_data in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    fill = 'DCE6F1' if ri % 2 == 0 else 'FFFFFF'
    for ci, text in enumerate(row_data):
        cell = row.cells[ci]
        cell.width = col_widths[ci]
        shade_cell(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        bold = (ci == 0)
        set_font(run, size=11, bold=bold)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1.5.2
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('1.5.2. Nhóm phát triển hệ thống')

add_body(
    'Nhóm này bao gồm Nhóm phát triển phần mềm, giữ vai trò đảm bảo vận hành, '
    'kiểm soát chất lượng và phát triển liên tục hệ thống trong suốt vòng đời dự án.',
    bold_parts=['Nhóm phát triển phần mềm']
)

add_body(
    'Nhóm phát triển hệ thống thực hiện phân tích yêu cầu, thiết kế kiến trúc, triển khai và '
    'bảo trì hệ thống trên nền tảng React.js (Frontend) và Node.js/Express (Backend), deployment '
    'lên Vercel và Railway/Render trong điều kiện ngân sách bằng 0 với free tier. Nhóm này '
    'tương tác với hệ thống thông qua quá trình phát triển, không phải qua giao diện người dùng cuối.',
    bold_parts=['Nhóm phát triển hệ thống']
)

# ══════════════════════════════════════════════════════════════════════════════
# 1.5.3
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('1.5.3. Nhóm hệ thống tích hợp và tự động hóa')

add_body(
    'Nhóm này bao gồm Hệ thống định thời (Timer) và Cổng thanh toán (Payment Gateway), đóng '
    'vai trò hỗ trợ xử lý nghiệp vụ tự động mà không có sự can thiệp trực tiếp của con người, '
    'đảm bảo tính liên tục và nhất quán của dữ liệu vận hành.',
    bold_parts=['Hệ thống định thời (Timer)', 'Cổng thanh toán (Payment Gateway)']
)

add_body(
    'Hệ thống định thời (Timer) là tác nhân phụ tự động kích hoạt các sự kiện định kỳ: reset '
    'streak khi người dùng không hoạt động quá 2 ngày, gia hạn gói subscription, kết thúc '
    'challenge khi hết thời hạn và gửi nhắc nhở trả thiết bị thuê. Timer đảm bảo tính nhất '
    'quán của dữ liệu gamification và hợp đồng thuê mà không cần xử lý thủ công.',
    bold_parts=['Hệ thống định thời (Timer)']
)

add_body(
    'Cổng thanh toán (Payment Gateway) là hệ thống bên ngoài (VNPay/Momo ở chế độ sandbox), '
    'tiếp nhận yêu cầu thanh toán từ FitFuel+, xử lý giao dịch và trả về kết quả qua callback '
    'URL. Mọi luồng thanh toán của nền tảng — đặt đồ ăn, mua/thuê gear, nạp FitCoin, gia hạn '
    'membership — đều đi qua tác nhân này.',
    bold_parts=['Cổng thanh toán (Payment Gateway)']
)

# ══════════════════════════════════════════════════════════════════════════════
# Summary table 1.5.2 & 1.5.3
# ══════════════════════════════════════════════════════════════════════════════
add_note('Bảng 1.5.2 – Tổng hợp nhóm phát triển và hệ thống tích hợp')

tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Tác nhân', 'Loại', 'Vai trò trong hệ thống']
rows2 = [
    ['Nhóm phát triển', 'Người dùng – Nội bộ','Phân tích, thiết kế, triển khai, bảo trì hệ thống'],
    ['Timer',           'Hệ thống – Phụ',     'Tự động reset streak, gia hạn subscription, kết thúc challenge'],
    ['Payment Gateway', 'Hệ thống – Phụ',     'Xử lý giao dịch thanh toán, trả kết quả qua callback URL'],
]

col_widths2 = [Cm(3.5), Cm(4.0), Cm(9.5)]
for i, cell in enumerate(tbl2.rows[0].cells):
    cell.width = col_widths2[i]
    shade_cell(cell, '1F497D')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(headers2[i])
    set_font(run, size=12, bold=True, color=(255, 255, 255))

for ri, row_data in enumerate(rows2):
    row = tbl2.rows[ri + 1]
    fill = 'DCE6F1' if ri % 2 == 0 else 'FFFFFF'
    for ci, text in enumerate(row_data):
        cell = row.cells[ci]
        cell.width = col_widths2[ci]
        shade_cell(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=11, bold=(ci == 0))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1.5.4
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('1.5.4. Mối quan hệ và ý nghĩa của các bên liên quan')

add_body(
    'Các nhóm người dùng trong FitFuel+ tồn tại mối quan hệ phụ thuộc lẫn nhau tạo nên giá '
    'trị hệ sinh thái. Member cung cấp dữ liệu gym log — dữ liệu này là đầu vào để AI gợi ý '
    'thức ăn từ Food Vendor và gợi ý thiết bị từ Gear Seller, tạo ra vòng lặp giá trị không '
    'thể có nếu các mảng hoạt động độc lập.',
    bold_parts=['Member', 'Food Vendor', 'Gear Seller']
)

add_body(
    'Food Vendor và Gear Seller cung cấp catalog sản phẩm, quyết định trực tiếp đến trải nghiệm '
    'mua sắm của Member và Guest. Gym Owner mang lại bối cảnh tập luyện, kết nối hoạt động gym '
    'trực tiếp với dữ liệu số trên nền tảng. Gym Owner đảm bảo tính tin cậy của toàn bộ hệ sinh '
    'thái thông qua kiểm duyệt và xử lý tranh chấp trên toàn nền tảng. Timer và Payment Gateway duy trì tính '
    'liên tục và an toàn tài chính cho mọi giao dịch.',
    bold_parts=['Food Vendor', 'Gear Seller', 'Gym Owner', 'Timer', 'Payment Gateway']
)

add_body(
    'Việc xác định đúng và đầy đủ các bên liên quan là cơ sở để thiết kế phân quyền truy cập, '
    'xây dựng use case, phân tích luồng dữ liệu và đảm bảo tính khả thi trong điều kiện phát '
    'triển 6 tuần với đội ngũ 5 thành viên của dự án FitFuel+.'
)

# ══════════════════════════════════════════════════════════════════════════════
# Relationship matrix table
# ══════════════════════════════════════════════════════════════════════════════
add_note('Bảng 1.5.3 – Ma trận quan hệ giữa các bên liên quan')

actors = ['Guest', 'Member', 'Food Vendor', 'Gear Seller', 'Gym Owner', 'Timer', 'Payment GW']
n = len(actors)
tbl3 = doc.add_table(rows=n+1, cols=n+1)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

relations = {
    ('Guest',       'Food Vendor'):   '🛒 Đặt',
    ('Guest',       'Gear Seller'):   '👀 Xem',
    ('Guest',       'Payment GW'):    '💳 Thanh toán',
    ('Member',      'Food Vendor'):   '🛒 Đặt + AI',
    ('Member',      'Gear Seller'):   '🛒 Thuê/Mua',
    ('Member',      'Gym Owner'):     '🏋 Check-in',
    ('Member',      'Payment GW'):    '💳 Thanh toán',
    ('Member',      'Timer'):         '⏱ Streak/Challenge',
    ('Food Vendor', 'Gym Owner'):     '✅ Được duyệt',
    ('Gear Seller', 'Gym Owner'):     '✅ Listing duyệt',
    ('Gym Owner',   'Timer'):         '⚙ Cấu hình',
    ('Gym Owner',   'Payment GW'):    '⚙ Giám sát',
}

# Header row
shade_cell(tbl3.rows[0].cells[0], '1F497D')
for ci, actor in enumerate(actors):
    cell = tbl3.rows[0].cells[ci+1]
    shade_cell(cell, '1F497D')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(actor)
    set_font(run, size=9, bold=True, color=(255,255,255))

for ri, row_actor in enumerate(actors):
    row = tbl3.rows[ri+1]
    # Row header
    hcell = row.cells[0]
    shade_cell(hcell, '1F497D')
    p = hcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(row_actor)
    set_font(run, size=9, bold=True, color=(255,255,255))

    for ci, col_actor in enumerate(actors):
        cell = row.cells[ci+1]
        if ri == ci:
            shade_cell(cell, 'BFBFBF')
        else:
            key = (row_actor, col_actor)
            val = relations.get(key, '')
            fill = 'E2EFDA' if val else 'FFFFFF'
            shade_cell(cell, fill)
            if val:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                set_font(run, size=8)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
output_path = r'D:\doanWEDKD\FitFuel_Stakeholder_Analysis.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
